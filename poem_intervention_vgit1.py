# -*- coding: utf-8 -*-
"""

poem intervention v2


Created on Sun Aug 12 11:19:42 2018

@author: Sylvan
"""

from Crypto.Cipher import AES
import base64
import os

import re
import docx
from docx import Document
import requests
from bs4 import BeautifulSoup
import random

text_list = []
def getText(text_list):
    pages = ['source docx files']
    for filename in pages:
        doc = Document(filename)
        fullText = []
        flattened = []
        for p in doc.paragraphs:
            fullText.append(p.text.split(' '))


        for sublist in fullText:
            for val in sublist:
                flattened.append(val)

        text_list.append(flattened)


getText(text_list)
#print(text_list)

#print(len(text_list))












def regex_1(doc_obj, regex_list_one, regex_list_two, regex_list_three):

    #main regex replacer, replaces words matching regex with random words from word generator
    pull_counter = 0
    counter = 0

##REGEX 1 - THE CENTRAL REPS
    for p in doc_obj.paragraphs:
        #for regex in regex_list_one:
        if regex_list_one.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex_list_one.search(inline[i].text):
                    r_word = RANDOM_WORD_INCREASER_2(pull_counter, counter)
                    text = regex_list_one.sub(r_word,inline[i].text)
                    inline[i].text = text
                    counter +=1
                    pull_counter +=1
                    if counter == 4:
                        counter = 0
##REGEX 2 - SLEEPING
    r_word_static = str(WORD_GEN(text_list[counter]))
    for p in doc_obj.paragraphs:
        #for regex in regex_list_one:
        if regex_list_two.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex_list_two.search(inline[i].text):
                    text = regex_list_two.sub(r_word_static,inline[i].text)
                    inline[i].text = text

###REGEX 3 - THE PHRASES
    for p in doc_obj.paragraphs:
        #for regex in regex_list_one:
        if regex_list_three.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex_list_three.search(inline[i].text):
                    r_word3 = RANDOM_WORD_INCREASER_3(random.randint(0,3))
                    text = regex_list_three.sub(r_word3,inline[i].text)
                    inline[i].text = text
   
    

#word generator, selects words at random from input files
def WORD_GEN(inputdoc):
    r_WORD_GEN = random.randint(0, len(inputdoc)-1)
    r_word  = inputdoc[r_WORD_GEN]
    if len(r_word) > 3:
        return r_word
    else:
        return WORD_GEN(inputdoc)


def WORD_GEN_COMPLEX(inputdoc, phrase_length):
    r_WORD_GEN = random.randint(0, len(inputdoc)-15)
    r_word  = inputdoc[r_WORD_GEN:r_WORD_GEN + phrase_length]
    r_word = (' ').join(r_word)
    return r_word

#WORD_GEN_COMPLEX(text_list[0], 6)
#word_gen(text_list[1])


def RANDOM_WORD_INCREASER(pull_counter, counter):
    pc = pull_counter
    rgen = str(WORD_GEN(text_list[counter]))
    rgen2 = str(WORD_GEN(text_list[counter]))
    if pc < 3:
        return rgen
    elif pc in range (3,8):
        return rgen + "/_/" + rgen
    elif pc in range (8,12):
        return rgen[0:(len(rgen)//2)] + "//" + rgen2 + "//" + rgen[(len(rgen)//2):len(rgen)]
    else:
        return rgen2[:(len(rgen2)//2)] + rgen[:(len(rgen)//2)] + "//" + rgen2 + "//" + rgen[(len(rgen)//2):]


def RANDOM_WORD_INCREASER_2(pull_counter, counter):
    pc = pull_counter
    rgen = str(WORD_GEN(text_list[counter]))
    rgen2 = str(WORD_GEN(text_list[counter]))
    if pc < 3:
        return rgen
    elif pc in range (3,8):
        return rgen + "/_/" + rgen
    elif pc in range (8,12):
        return rgen[0:(len(rgen)//2)] + "/_/" + rgen2 + "/_/" + rgen[(len(rgen)//2):len(rgen)]
    elif pc in range (13,17):
        return rgen
    elif pc in range(18,25):
        return rgen2[:(len(rgen2)//2)] + rgen2 + rgen[(len(rgen)//2):]
    else:
        rgen2 + " " + rgen


def RANDOM_WORD_INCREASER_3(counter):
    random_lengths = [2,2,6,6,6,6,6,10,10,10]
    chooser = random.randint(0,len(random_lengths)-1)
    phrase_length = random_lengths[chooser]
    rgen = WORD_GEN_COMPLEX(text_list[counter],phrase_length)
   # rgen = (" ").join(rgen)
    return rgen



RANDOM_WORD_INCREASER_3(3)

#RANDOM_WORD_INCREASER(17,2)



def conjunction_randomizer(doc_obj):
    #conjunction randomizer
    conj_list_regex = [re.compile(r" yet "), re.compile(r'but'), re.compile(r' and '), re.compile(r' or ')]
    conj_list = ['yet','and','but','or']
    for p in doc_obj.paragraphs:
        for word in conj_list_regex:
            if word.match(p.text):
                inline = p.runs
                replace_chance = random.randint(0, 99)
                if replace_chance > 49:
                    for i in range(len(inline)):
                        if word.match(inline[i].text):
                            r_list_gen = random.randint(0,len(conj_list)-1)
                            r_conj = conj_list[r_list_gen]
                            text = word.sub(r_conj, inline[i].text)
                            inline[i].text = text




def encryption(word):
    BLOCK_SIZE = 16
    PADDING = '{'

    pad = lambda s: s + (BLOCK_SIZE - len(s) % BLOCK_SIZE) * PADDING

    EncodeAES = lambda c, s: base64.b64encode(c.encrypt(pad(s)))

    secret = os.urandom(BLOCK_SIZE)

    #print('encryption key:', secret)

    cipher = AES.new(secret)

    encoded = EncodeAES(cipher, word)
    encoded_str = str(encoded)
    encoded_str = encoded_str[2:-4]
    return encoded_str


def short_encrypter(doc_obj):
        for p in doc_obj.paragraphs:
            inline = p.runs
            replace_chance = random.randint(1, 99)
            if replace_chance > 27:
                for i in range(len(inline)):
                    if len(inline[i].text) in range(1,2):
                        print(inline[i].text)
                        text = encryption(inline[i].text)
                        inline[i].text = text
                        print(inline[i].text)

#short_encrypter(doc)



################### master func #################

def save_func(doc_obj,save_counter):
    doc.save('result'+str(save_counter)+'.docx')


save_counter = 40


def docx_replace_regex(doc_obj, regex_list_one, regex_list_two,regex_list_three,save_counter):

    #while save_counter < 25:
    #main regex
    regex_1(doc_obj, regex_list_one, regex_list_two,regex_list_three)
    conjunction_randomizer(doc_obj)
    #short_encrypter(doc_obj)
    save_func(doc_obj, save_counter)
        #save_counter +=1



##############################################

#regex_list_one = re.compile(r'\b(?:%s)\b' % '|'.join(word_list_one))
regex_list_one = re.compile('words|one',re.IGNORECASE)
regex_list_two = re.compile('words|two',re.IGNORECASE)
regex_list_three = re.compile("words three",re.IGNORECASE)



filename = "source_file.docx"
doc = Document(filename)

docx_replace_regex(doc, regex_list_one, regex_list_two, regex_list_three, save_counter)
